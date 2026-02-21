import docx
import os
import json
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TEXT = WORD_NAMESPACE + "t"

def get_accepted_text(p):
    """Return text of a paragraph after accepting all changes"""
    xml = p._p.xml
    if "w:del" in xml or "w:ins" in xml:
        tree = XML(xml)
        runs = (node.text for node in tree.iter(TEXT) if node.text)
        return "".join(runs)
    else:
        return p.text
    
def get_accepted_text_cell(c):
    return "".join([get_accepted_text(p) for p in c.paragraphs])

def extract_hpo_from_reports(doc_name='./WES20250001-140/.docx'):
    document = Document(doc_name)
    t = document.tables[0]
    description = get_accepted_text_cell(t.cell(3, 0))
    
    temp, index = 0, 6
    for row in t.rows:
        temp += 1
    # >1 tables concatenaed
    if temp == 4:
        index = 2
        t = document.tables[1]
        

    results = []
    for i in range(index, 100):
        if t.cell(i, 0).text.startswith('分析结果：'):
            break
        temp = []
        for cell in t.rows[i].cells:
            if not temp or get_accepted_text_cell(cell) != temp[-1]:
                temp.append(get_accepted_text_cell(cell))
        results.append([doc_name]+temp)
        
    return description, results

if __name__ == '__main__':
    tgt_folder = './WES20250001-140/'
    samples = []
    samples_no_info = []
    for file in os.listdir(tgt_folder):
        if file.endswith('.docx') and not file.startswith('~'):
            description, results = extract_hpo_from_reports(os.path.join(tgt_folder, file))
            out_dict = {
                'description': description,
                'results': results
            }
            out_dict_no_info = {
                'description': description,
                'results': list(map(lambda line: line[4:], results))
            }
            samples.append(json.dumps(out_dict, ensure_ascii=False)+'\n')
            samples_no_info.append(json.dumps(out_dict_no_info, ensure_ascii=False)+'\n')
    with open('./processed_data.json', 'w', encoding='utf-8') as f:
        f.writelines(samples)
    with open('./processed_data_no_info.json', 'w', encoding='utf-8') as f:
        f.writelines(samples_no_info)

# 表示xx和患者的关系

# 不良孕史作为表型                                妥协
# 299. 甲基丙二酸血症?                            孩子的表型
# 332. 近视+重度近视                              只保留子表型，子+父表型不算错
# 335. 视网膜脱离/轻度近视 ？                      从轻度近视发展到重度近视 / 视网膜脱离漏了
# 383. 输卵管异常                                 前后病例变化导致
# 439. 原位癌可能？/病例报告/疑似xx                 即使前面有疑问，也应当被识别
# 没有xxxxx/否认xxxxx                             否认表型不应当被提取 

# 1000 中文文献，国家罕见病病例交流讨论会/中华xxxx杂志/全文xxx数据库
# 1000 测序报告
# 1000 住院志